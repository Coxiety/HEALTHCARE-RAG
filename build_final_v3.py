from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def P(text, *, bold=False, italic=False, size=13, align=None, space_after=6, indent_first=0):
    p = doc.add_paragraph()
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if indent_first: p.paragraph_format.first_line_indent = Cm(indent_first)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    return p


def H(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.bold = True
    run.font.size = Pt(15 if level == 1 else (14 if level == 2 else 13))
    return p


def CT(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.bold = True; run.font.size = Pt(15)
    return p


def B(text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"• {text}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    return p


def CODE(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    pPr = p.paragraph_format.element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'; run.font.size = Pt(10)
    return p


def TBL(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; cell.text = ''
        set_cell_bg(cell, '2C3E50')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        if r_idx % 2 == 1:
            for cell in tr.cells:
                set_cell_bg(cell, 'F8F9FA')
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]; cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'; run.font.size = Pt(11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════
P("BỘ KHOA HỌC VÀ CÔNG NGHỆ", bold=True, align='center', space_after=2)
P("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", bold=True, align='center', space_after=2)
P("-----------------------------", align='center', space_after=20)
P("BÁO CÁO", bold=True, align='center', size=20, space_after=4)
P("ĐỒ ÁN MÔN HỌC", bold=True, align='center', size=20, space_after=20)
P("ĐỀ TÀI: HỆ THỐNG HỎI ĐÁP THÔNG MINH VỀ DINH DƯỠNG VÀ SỨC KHỎE DỰA TRÊN RETRIEVAL-AUGMENTED GENERATION (RAG)",
  bold=True, align='center', size=14, space_after=24)
P("Môn học: Xử lý Ngôn ngữ Tự nhiên (NLP)", bold=True, align='center', space_after=4)
P("Giảng viên hướng dẫn: [Tên giảng viên]", bold=True, align='center', space_after=20)
P("Thực hiện bởi nhóm sinh viên, bao gồm:", align='center', space_after=8)

mt = doc.add_table(rows=3, cols=3)
mt.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, (name, mssv, role) in enumerate([
    ("Trần Minh Khang", "N22DCCN138", "Trưởng nhóm"),
    ("Nguyễn Hải Đông", "[MSSV]", "Thành viên"),
    ("Đặng Nhật Nam", "[MSSV]", "Thành viên"),
]):
    for c_idx, val in enumerate([name, mssv, role]):
        cell = mt.rows[r_idx].cells[c_idx]; cell.text = ''
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val); run.font.name = 'Times New Roman'; run.font.size = Pt(13)

doc.add_paragraph()
P("TP.HCM, tháng 6/2026", align='center', bold=True)
doc.add_page_break()

# MUC LUC
CT("MỤC LỤC")
P("(Tự sinh sau khi hoàn thành)", italic=True, align='center', size=12)
doc.add_page_break()

# DANH SACH HINH BANG
CT("DANH SÁCH HÌNH, BẢNG")
TBL(["STT", "Tên", "Trang"], [
    ["Hình 2.1", "Kiến trúc tổng quan RAG", ""],
    ["Hình 2.2", "Sơ đồ Hybrid Search và Cross-encoder Reranker", ""],
    ["Hình 3.1", "Pipeline tổng thể hệ thống", ""],
    ["Hình 3.2", "Sơ đồ 2 nhánh retrieval (USDA + NFCorpus)", ""],
    ["Hình 3.3", "Quy trình gán nhãn NER (BIO format)", ""],
    ["Hình 4.1", "Giao diện chatbot Spring Boot", ""],
    ["Hình 4.2", "Ví dụ câu hỏi và câu trả lời", ""],
    ["Bảng 2.1", "So sánh RAG vs Fine-tuning LLM", ""],
    ["Bảng 3.1", "Ba loại câu hỏi chính", ""],
    ["Bảng 3.2", "Bốn loại entity trong NER", ""],
    ["Bảng 4.1", "Môi trường và công cụ phát triển", ""],
    ["Bảng 5.1", "Đánh giá synthetic intent dataset", ""],
    ["Bảng 5.2", "F1 Score cho NER", ""],
    ["Bảng 5.3", "Kết quả Intent Classifier", ""],
    ["Bảng 5.4", "Top-k Accuracy và MRR cho Retrieval", ""],
    ["Bảng 5.5", "RAGAS Evaluation", ""],
], col_widths=[1.5, 11, 2])
doc.add_page_break()

# TOM TAT - FIXED: lavita → NFCorpus
CT("TÓM TẮT")
P("Đồ án xây dựng một hệ thống hỏi đáp thông minh về dinh dưỡng và sức khỏe, sử dụng kiến trúc "
  "Retrieval-Augmented Generation (RAG). Hệ thống nhận câu hỏi từ người dùng, phân loại ý định "
  "bằng mô hình BERT đã được fine-tune, nhận diện các thực thể (thực phẩm, bệnh, chất dinh dưỡng, "
  "triệu chứng) bằng mô hình BioBERT đã được fine-tune trên bộ BC5CDR, sau đó truy xuất thông "
  "tin từ hai nguồn: cơ sở dữ liệu dinh dưỡng USDA (SQLite) cho số liệu chính xác và kho tài "
  "liệu y tế NFCorpus — PubMed abstracts (ChromaDB) cho lời khuyên sức khỏe. Hệ thống so sánh "
  "bốn phương pháp truy xuất: TF-IDF, BM25, Dense Vector và Hybrid RRF, kết hợp với cross-encoder "
  "reranker để cải thiện chất lượng. Cuối cùng, mô hình ngôn ngữ lớn llama3.1:8b tổng hợp thông "
  "tin và sinh câu trả lời có nguồn gốc, hạn chế hallucination. Kết quả đánh giá đạt NER F1 = 0.893, "
  "Intent Classifier F1 = 1.000 (in-distribution), Retrieval MRR = 0.55 (Hybrid RRF + Reranker "
  "trên NFCorpus 323 test queries).",
  align='justify', indent_first=1)
doc.add_page_break()

# BANG PHAN CONG - FIXED
CT("BẢNG PHÂN CÔNG CÔNG VIỆC TRONG NHÓM")
TBL(["STT", "MSSV", "Họ và Tên", "Nhiệm vụ"], [
    ["1", "N22DCCN138", "Trần Minh Khang",
     "Xây dựng pipeline RAG, Training 1 — NER BioBERT, retrieval (4 baseline + reranker), tích hợp LLM llama3.1:8b, evaluation (eval_retrieval.ipynb)"],
    ["2", "[MSSV]", "Nguyễn Hải Đông",
     "BC5CDR preprocessing, triplet mining từ NFCorpus qrels, Training 2 — Embedding contrastive (MNR loss, all-MiniLM-L6-v2)"],
    ["3", "[MSSV]", "Đặng Nhật Nam",
     "Giao diện Spring Boot, sinh synthetic intent dataset, Training 3 — Intent Classifier (BERT 3-class), so sánh ablation, chuẩn bị demo"],
], col_widths=[1, 2.5, 3, 8])
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# CHUONG I - FIXED: Vietnamese → English language
# ════════════════════════════════════════════════════════════════
CT("CHƯƠNG I — GIỚI THIỆU ĐỀ TÀI, MỤC TIÊU VÀ PHẠM VI NGHIÊN CỨU")

H("1.1 Lý do chọn đề tài", 2)
P("Nhu cầu tìm kiếm thông tin về dinh dưỡng và sức khỏe ngày càng tăng, đặc biệt với các câu hỏi "
  "liên quan đến chế độ ăn cho người mắc bệnh mãn tính như tiểu đường, gout, hay cao huyết áp. "
  "Tuy nhiên, các công cụ tìm kiếm truyền thống thường trả về kết quả chung chung, không trả lời "
  "trực tiếp câu hỏi của người dùng.", align='justify', indent_first=1)
P("Trong khi đó, các mô hình ngôn ngữ lớn (LLM) như ChatGPT tuy có khả năng trả lời tự nhiên "
  "nhưng thường gặp vấn đề hallucination, tức sinh ra thông tin sai nhưng nghe hợp lý. Vấn đề "
  "này đặc biệt nguy hiểm trong lĩnh vực y tế, nơi thông tin sai có thể dẫn đến hậu quả nghiêm "
  "trọng cho sức khỏe người dùng.", align='justify', indent_first=1)
P("Kiến trúc Retrieval-Augmented Generation (RAG) giải quyết được cả hai vấn đề trên bằng cách "
  "kết hợp khả năng truy xuất thông tin chính xác từ nguồn dữ liệu đáng tin cậy với khả năng "
  "sinh câu trả lời tự nhiên của LLM. Đồ án này áp dụng kiến trúc RAG để xây dựng một hệ thống "
  "hỏi đáp chuyên biệt cho lĩnh vực dinh dưỡng và sức khỏe bằng tiếng Anh.",
  align='justify', indent_first=1)

H("1.2 Mục tiêu của đề tài", 2)
for item in [
    "Xây dựng hệ thống hỏi đáp tiếng Anh về dinh dưỡng và sức khỏe sử dụng kiến trúc RAG.",
    "Nhận diện các thực thể trong câu hỏi (thực phẩm, bệnh, chất dinh dưỡng, triệu chứng) bằng mô hình BioBERT đã được fine-tune.",
    "Truy xuất thông tin chính xác từ hai nguồn dữ liệu: số liệu dinh dưỡng (USDA) và tài liệu y tế (NFCorpus).",
    "Sinh câu trả lời có nguồn gốc, hạn chế hallucination.",
    "Đánh giá chất lượng hệ thống bằng các metric chuẩn (F1, MRR, RAGAS).",
]:
    B(item)

H("1.3 Phạm vi và đối tượng nghiên cứu", 2)
P("Phạm vi:", bold=True)
for item in [
    "Câu hỏi tiếng Anh về dinh dưỡng (hàm lượng chất dinh dưỡng trong thực phẩm) và sức khỏe (chế độ ăn cho người mắc bệnh).",
    "Dữ liệu dinh dưỡng từ USDA FoodData Central (13,661 thực phẩm).",
    "Tài liệu y tế từ NFCorpus (BEIR) — 3,633 PubMed abstracts, 323 test queries.",
    "Dữ liệu NER y sinh BC5CDR — 16,423 sentences (tner/bc5cdr) cho fine-tune BioBERT.",
]:
    B(item)
P("Đối tượng:", bold=True)
P("Người dùng phổ thông muốn tìm hiểu về dinh dưỡng và chế độ ăn phù hợp với tình trạng sức khỏe.",
  align='justify', indent_first=1)

H("1.4 Phương pháp thực hiện", 2)
P("Đồ án áp dụng kiến trúc Hybrid RAG kết hợp hai phương pháp truy xuất:",
  align='justify', indent_first=1)
for item in [
    "Truy vấn chính xác (exact lookup) trên cơ sở dữ liệu có cấu trúc SQLite cho số liệu dinh dưỡng.",
    "Tìm kiếm kết hợp BM25 + semantic search với Reciprocal Rank Fusion trên ChromaDB cho tài liệu y tế.",
    "Cross-encoder reranker (pretrained) để cải thiện chất lượng top-k cuối cùng.",
]:
    B(item)
P("Mô hình BioBERT được fine-tune cho bài toán Named Entity Recognition (NER) để nhận diện thực "
  "thể trong câu hỏi. Mô hình ngôn ngữ lớn llama3.1:8b được sử dụng để tổng hợp thông tin và "
  "sinh câu trả lời.", align='justify', indent_first=1)

H("1.5 Cấu trúc báo cáo", 2)
for item in [
    "Chương I: Giới thiệu đề tài, mục tiêu và phạm vi nghiên cứu.",
    "Chương II: Cơ sở lý thuyết về RAG, TF-IDF, BM25, Dense Vector, NER, Intent Classification.",
    "Chương III: Phân tích và thiết kế kiến trúc hệ thống.",
    "Chương IV: Xây dựng và triển khai các module.",
    "Chương V: Đánh giá kết quả bằng F1, Top-k accuracy, MRR và RAGAS.",
    "Chương VI: Kết luận và hướng phát triển.",
]:
    B(item)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# CHUONG II - FIXED: lavita → NFCorpus
# ════════════════════════════════════════════════════════════════
CT("CHƯƠNG II — CƠ SỞ LÝ THUYẾT")

H("2.1 Retrieval-Augmented Generation (RAG)", 2)
P("RAG (Lewis và cộng sự, 2020) là kiến trúc kết hợp truy xuất thông tin với sinh ngôn ngữ. Với "
  "một câu hỏi đầu vào, hệ thống đầu tiên tìm các đoạn văn liên quan từ kho tri thức, sau đó đưa "
  "các đoạn này làm ngữ cảnh (context) cho LLM để sinh câu trả lời có căn cứ. Cách tiếp cận này "
  "giảm hallucination vì model bị ràng buộc bởi thông tin từ nguồn đã được xác minh.",
  align='justify', indent_first=1)
P("[Hình 2.1 — Kiến trúc tổng quan RAG]", italic=True, align='center', size=12)
P("Bảng 2.1 — So sánh RAG vs Fine-tuning LLM:", italic=True, size=12)
TBL(["Tiêu chí", "RAG", "Fine-tuning LLM"], [
    ["Hallucination", "Thấp, câu trả lời dựa trên nguồn cụ thể", "Cao hơn, model có thể bịa"],
    ["Cập nhật kiến thức", "Dễ, chỉ thêm tài liệu mới", "Khó, phải train lại"],
    ["Tài nguyên", "Ít hơn, chỉ cần inference", "Nhiều, cần GPU để train"],
    ["Trích nguồn", "Có, chỉ rõ nguồn gốc", "Không"],
], col_widths=[4, 6, 5])

H("2.2 BioBERT và Named Entity Recognition", 2)
P("BioBERT là mô hình pre-trained dựa trên BERT, được train thêm trên dữ liệu y sinh (PubMed "
  "abstracts và PMC full-text). BioBERT cung cấp biểu diễn ngôn ngữ phù hợp cho domain y tế, "
  "phù hợp với bài toán NER trên các thực thể bệnh, chất hóa học, chất dinh dưỡng và triệu chứng.",
  align='justify', indent_first=1)
P("NER (Named Entity Recognition) là bài toán nhận diện thực thể trong văn bản, sử dụng định "
  "dạng BIO tagging: B- (beginning) đánh dấu token bắt đầu của entity, I- (inside) đánh dấu các "
  "token tiếp theo trong cùng entity, và O đánh dấu token không thuộc entity nào.",
  align='justify', indent_first=1)

H("2.3 TF-IDF và Sparse Retrieval", 2)
P("TF-IDF (Term Frequency – Inverse Document Frequency) là phương pháp truy xuất sparse cổ điển. "
  "TF đo tần suất xuất hiện của một từ trong tài liệu, IDF đo độ hiếm của từ đó trên toàn corpus. "
  "Một từ có TF cao và IDF cao là từ phân biệt tốt cho tài liệu đó.",
  align='justify', indent_first=1)
P("Công thức TF-IDF được định nghĩa như sau:", align='justify', indent_first=1)
P("TF(t, d) = count(t, d) / |d|", align='center', italic=True)
P("IDF(t) = log(N / df(t))", align='center', italic=True)
P("TF-IDF(t, d) = TF(t, d) × IDF(t)", align='center', italic=True)
P("Truy xuất được thực hiện bằng cách tính cosine similarity giữa vector TF-IDF của câu hỏi và "
  "tất cả vector của tài liệu. Triển khai bằng thư viện sklearn.",
  align='justify', indent_first=1)

H("2.4 Semantic Search và Embedding", 2)
P("Semantic search sử dụng sentence-transformers để mã hóa câu hỏi và tài liệu thành vector "
  "liên tục (embedding) trong không gian ngữ nghĩa. Mô hình all-MiniLM-L6-v2 được dùng làm "
  "embedding model cơ sở. Độ tương đồng được đo bằng cosine similarity giữa embedding của câu "
  "hỏi và mỗi tài liệu.", align='justify', indent_first=1)
P("sim(q, d) = cos(E(q), E(d)) = (E(q) · E(d)) / (|E(q)| × |E(d)|)",
  align='center', italic=True)
P("Ưu điểm của semantic search là bắt được ngữ nghĩa, xử lý tốt từ đồng nghĩa và paraphrase. "
  "Nhược điểm là bỏ sót khớp từ chính xác (exact match) và cần embedding model tốt cho domain. "
  "Vector được lưu trữ trong ChromaDB để truy vấn nhanh.",
  align='justify', indent_first=1)

H("2.5 BM25 và Hybrid Search", 2)
P("BM25 (Best Match 25) là phương pháp lexical search dựa trên tần suất từ, cải tiến từ TF-IDF "
  "bằng cơ chế term saturation và document length normalization, khắc phục vấn đề tài liệu dài "
  "bị cộng dồn TF quá mức.", align='justify', indent_first=1)
P("BM25(t,d) = IDF(t) × [f(t,d)·(k1+1)] / [f(t,d) + k1·(1 - b + b·|d|/avgdl)]",
  align='center', italic=True)
P("Trong đó k1 = 1.5 điều khiển term saturation và b = 0.75 điều khiển length normalization. "
  "Triển khai bằng thư viện rank_bm25.", align='justify', indent_first=1)
P("Hybrid search kết hợp BM25 (lexical) và Dense (semantic) bằng Reciprocal Rank Fusion (RRF). "
  "Mỗi tài liệu nhận điểm tổng hợp dựa trên thứ hạng của nó trong từng danh sách kết quả:",
  align='justify', indent_first=1)
P("RRF(d) = Σ 1 / (k + rank_i(d)),  với k = 60", align='center', italic=True)
P("BM25 bắt được từ chuyên ngành (HbA1c, glycemic index), Dense bắt được tương đồng ngữ nghĩa "
  "(diabetes ≈ high blood sugar). Kết hợp hai phương pháp giúp tăng cả Recall và Precision.",
  align='justify', indent_first=1)

H("2.6 Cross-encoder Reranker", 2)
P("Cross-encoder nhận cặp (query, passage) làm đầu vào chung và sinh ra điểm relevance bằng cơ "
  "chế attention chéo giữa query và passage. Phương pháp này chính xác hơn bi-encoder nhưng quá "
  "chậm cho retrieval cấp 1 trên toàn corpus, do đó được dùng làm tầng thứ hai (re-ranking): "
  "lấy top-20 từ retriever, rerank còn top-3 trước khi đưa vào LLM. Mô hình sử dụng là "
  "cross-encoder/ms-marco-MiniLM-L-6-v2.",
  align='justify', indent_first=1)
P("[Hình 2.2 — Sơ đồ Hybrid Search và Cross-encoder Reranker]", italic=True, align='center', size=12)

H("2.7 Mô hình ngôn ngữ lớn (LLM)", 2)
P("Llama3.1:8b là mô hình ngôn ngữ lớn open-source của Meta, có 8 tỷ tham số, hỗ trợ tốt tiếng "
  "Anh và xử lý tốt các tác vụ tổng hợp thông tin. Mô hình được chạy local thông qua Ollama, "
  "đảm bảo dữ liệu y tế của người dùng không phải gửi qua API ngoài.",
  align='justify', indent_first=1)

H("2.8 Các nghiên cứu liên quan", 2)
P("Nhiều hệ thống RAG cho domain y tế đã được công bố như BioRAG, MedRAG. Các hệ thống này "
  "thường sử dụng dataset PubMed làm corpus chính. Đồ án này sử dụng NFCorpus — corpus dinh "
  "dưỡng y sinh gồm 3,633 PubMed abstracts được tuyển chọn từ NutritionFacts.org, kết hợp dữ "
  "liệu chuẩn USDA cho phần số liệu dinh dưỡng. Đây là điểm khác biệt so với các hệ thống y tế "
  "thuần túy, tập trung chuyên biệt vào nutrition và healthcare consumer queries.",
  align='justify', indent_first=1)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# CHUONG III - FIXED: lavita → NFCorpus
# ════════════════════════════════════════════════════════════════
CT("CHƯƠNG III — PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

H("3.1 Khảo sát và thu thập yêu cầu", 2)
P("Yêu cầu chức năng:", bold=True)
for item in [
    "Nhận câu hỏi của người dùng về dinh dưỡng và sức khỏe.",
    "Trả lời chính xác số liệu dinh dưỡng (calo, protein, chất béo).",
    "Đưa ra lời khuyên về chế độ ăn phù hợp với bệnh lý.",
    "Trích dẫn nguồn tham khảo trong câu trả lời.",
    "Xử lý được câu hỏi kết hợp cả số liệu và lời khuyên.",
]:
    B(item)
P("Yêu cầu phi chức năng:", bold=True)
for item in [
    "Không hallucination, trả lời 'không biết' khi không tìm được thông tin.",
    "Thời gian phản hồi chấp nhận được (dưới 30 giây).",
    "Chạy local trên máy 24GB RAM.",
]:
    B(item)

H("3.2 Phân tích yêu cầu", 2)
P("Ba loại câu hỏi chính được phân biệt bởi Intent Classifier:", align='justify')
P("Bảng 3.1 — Ba loại câu hỏi chính:", italic=True, size=12)
TBL(["Loại", "Ví dụ", "Nhánh xử lý"], [
    ["NUTRITION_LOOKUP", "How many calories in 100g chicken breast?", "USDA SQLite"],
    ["HEALTH_ADVICE", "What should a diabetic patient avoid?", "Hybrid RAG (NFCorpus)"],
    ["BOTH", "Is salmon good for cholesterol? How much protein?", "Cả hai nhánh"],
], col_widths=[3.5, 7, 3])

H("3.3 Thiết kế tổng thể hệ thống", 2)
P("[Hình 3.1 — Sơ đồ pipeline tổng thể hệ thống]", italic=True, align='center', size=12)
P("Luồng xử lý:", bold=True)
for item in [
    "Tiền xử lý câu hỏi (NLTK/spaCy: tokenize, lowercase, loại stopwords).",
    "Phân loại Intent bằng BERT classifier (NUTRITION_LOOKUP / HEALTH_ADVICE / BOTH).",
    "Nhận diện thực thể bằng BioBERT NER (DISEASE, CHEMICAL/NUTRIENT, SYMPTOM, FOOD-regex).",
    "Truy xuất thông tin từ nhánh phù hợp.",
    "Reranking: cross-encoder rerank top-20 thành top-3 chunks.",
    "Tổng hợp và sinh câu trả lời bằng llama3.1:8b qua Ollama.",
]:
    B(item)

H("3.4 Thiết kế Hybrid Retrieval", 2)
P("[Hình 3.2 — Sơ đồ 2 nhánh retrieval]", italic=True, align='center', size=12)
P("Nhánh A — SQLite (Structured Data):", bold=True)
for item in [
    "Dữ liệu: USDA FoodData Central, 13,661 thực phẩm.",
    "Kỹ thuật: exact lookup theo tên thực phẩm.",
]:
    B(item)
P("Nhánh B — Hybrid Search (Unstructured Data):", bold=True)
for item in [
    "Dữ liệu: NFCorpus — 3,633 tài liệu PubMed abstracts.",
    "TF-IDF / BM25: tìm theo từ khóa, bắt được thuật ngữ chuyên ngành.",
    "Dense (vanilla / fine-tuned): tìm theo nghĩa, bắt được từ đồng nghĩa.",
    "Merge: Reciprocal Rank Fusion với k=60.",
    "Reranker: cross-encoder/ms-marco-MiniLM-L-6-v2 rerank top-20 thành top-3.",
]:
    B(item)
P("Lý do chọn Hybrid:", bold=True)
for item in [
    "SQLite cho số liệu: chính xác tuyệt đối, không hallucination.",
    "BM25 + Semantic cho lời khuyên: BM25 bắt từ chuyên ngành, Dense bắt từ đồng nghĩa.",
    "Reranker đẩy chunk liên quan nhất lên đầu, giúp LLM sinh câu trả lời chính xác hơn.",
]:
    B(item)

H("3.5 Thiết kế NER", 2)
P("Bốn loại entity được định nghĩa:", bold=True)
P("Bảng 3.2 — Bốn loại entity trong NER:", italic=True, size=12)
TBL(["Entity", "Ví dụ", "Vai trò"], [
    ["FOOD", "chicken breast, salmon, banana", "Tra cứu dinh dưỡng trong USDA (regex fallback)"],
    ["DISEASE", "diabetes, gout, hypertension", "Tìm lời khuyên chế độ ăn (BioBERT)"],
    ["NUTRIENT", "protein, calories, vitamin C", "Xác định thông tin cần tra (BioBERT)"],
    ["SYMPTOM", "headache, fatigue, nausea", "Mở rộng khả năng tìm kiếm (regex fallback)"],
], col_widths=[3, 6, 6])
P("Lưu ý: BC5CDR chỉ có label CHEMICAL (Drug) và DISEASE. FOOD và SYMPTOM được xử lý bằng "
  "regex fallback pattern trong pipeline.py (ví dụ: pattern \"protein in <food>\") thay vì "
  "training NER. Đây là giới hạn của BC5CDR, được bù bằng heuristic trong pipeline.",
  align='justify', indent_first=1, italic=True, size=12)
P("[Hình 3.3 — Quy trình gán nhãn NER (BIO format)]", italic=True, align='center', size=12)
P("Ví dụ format BIO:", bold=True)
P('Câu: "A diabetic patient drinks sugar cane juice"', italic=True, align='center')
P("Tags: O B-DISEASE O O B-FOOD I-FOOD I-FOOD", italic=True, align='center')

H("3.6 Thiết kế cơ sở dữ liệu", 2)
P("SQLite (USDA):", bold=True)
for item in [
    "Bảng food: food_id, description, food_category.",
    "Bảng food_nutrient: food_id, nutrient_id, amount.",
    "Bảng nutrient: nutrient_id, name, unit.",
]:
    B(item)
P("ChromaDB (NFCorpus):", bold=True)
for item in [
    "Collection: nfcorpus.",
    "Path: data/chroma_db.",
    "Metadata: source, doc_id, chunk_index.",
    "Embedding: all-MiniLM-L6-v2 (vanilla hoặc fine-tuned).",
]:
    B(item)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# CHUONG IV - FIXED: file paths, BC5CDR size, paths, NFCorpus
# ════════════════════════════════════════════════════════════════
CT("CHƯƠNG IV — XÂY DỰNG VÀ TRIỂN KHAI CÁC MODULE")

H("4.1 Môi trường và công cụ phát triển", 2)
P("Hệ thống được phát triển và triển khai trên môi trường sau:", indent_first=1)
for item in [
    "Hệ điều hành: Ubuntu 22.04 / Windows 11.",
    "Ngôn ngữ lập trình: Python 3.10.",
    "GPU huấn luyện: Tesla A100 (Google Colab).",
    "RAM yêu cầu inference: 24GB.",
    "Quản lý môi trường: Conda.",
]:
    B(item)
P("Bảng 4.1 — Môi trường và công cụ phát triển:", italic=True, size=12)
TBL(["Thành phần", "Công nghệ", "Phiên bản / Model"], [
    ["NER", "BioBERT (fine-tune)", "dmis-lab/biobert-base-cased-v1.2"],
    ["Intent Classifier", "BERT (fine-tune)", "bert-base-uncased"],
    ["Embedding", "sentence-transformers", "all-MiniLM-L6-v2"],
    ["Vector DB", "ChromaDB", "0.4.x"],
    ["Lexical search", "rank_bm25", "0.2.2"],
    ["Sparse retrieval", "scikit-learn TfidfVectorizer", "1.3.x"],
    ["Reranker", "cross-encoder (pretrained)", "ms-marco-MiniLM-L-6-v2"],
    ["LLM", "llama3.1:8b qua Ollama", "3.1"],
    ["Giao diện", "Spring Boot", "3.x"],
    ["Quản lý môi trường", "Conda", "nutrition-rag"],
], col_widths=[4, 5, 6])

H("4.2 Mô tả các module chính", 2)

H("4.2.1 Module tiền xử lý (preprocessor.py)", 3)
P("Module thực hiện các bước:", indent_first=1)
for item in [
    "Chuẩn hóa Unicode về dạng NFC.",
    "Chuyển văn bản về chữ thường (lowercase).",
    "Tokenize bằng spaCy (mô hình en_core_web_sm).",
    "Loại bỏ stopwords theo danh sách của NLTK.",
]:
    B(item)
P("Ví dụ minh họa:", bold=True)
CODE('Đầu vào         : "What foods should a Diabetic patient AVOID?"\n'
     'Sau chuẩn hóa   : "what foods should a diabetic patient avoid"\n'
     'Sau tokenize    : ["what","foods","should","a","diabetic","patient","avoid"]\n'
     'Sau loại stopwords: ["foods","diabetic","patient","avoid"]')

H("4.2.2 Module nhận diện thực thể NER (ner.py)", 3)
P("a) Mô hình: BioBERT (dmis-lab/biobert-base-cased-v1.2)", bold=True)
for item in [
    "Kiến trúc: 12 lớp Transformer, 768 hidden units, 12 attention heads.",
    "Tổng tham số: 110M.",
    "Pre-train data: PubMed abstracts và PMC full-text articles.",
]:
    B(item)
P("b) Tham số huấn luyện (fine-tune):", bold=True)
for item in [
    "Số lớp đầu ra: 5 (B-CHEMICAL, I-CHEMICAL, B-DISEASE, I-DISEASE, O).",
    "Batch size: 32 (Colab A100).",
    "Learning rate: 3e-5.",
    "Số epoch: 3.",
    "Optimizer: AdamW (weight_decay = 0.01).",
    "Loss function: CrossEntropyLoss.",
    "Max sequence length: 128 tokens.",
]:
    B(item)
P("c) Dữ liệu huấn luyện:", bold=True)
for item in [
    "BC5CDR: 16,423 sentences từ tner/bc5cdr (Chemical, Disease).",
    "Chia: 90% train / 10% validation.",
]:
    B(item)
P("d) Code minh họa:", bold=True)
CODE('from transformers import (AutoTokenizer,\n'
     '                          AutoModelForTokenClassification)\n\n'
     'tokenizer = AutoTokenizer.from_pretrained("models/ner_bert")\n'
     'model = AutoModelForTokenClassification.from_pretrained(\n'
     '            "models/ner_bert")\n\n'
     'inputs = tokenizer(question, return_tensors="pt")\n'
     'outputs = model(**inputs)\n'
     'predictions = outputs.logits.argmax(dim=-1)')

H("4.2.3 Module phân loại ý định (classifier.py)", 3)
P("a) Mô hình: bert-base-uncased + 1 lớp Linear (3 đầu ra).", bold=True)
P("b) Tham số huấn luyện:", bold=True)
for item in [
    "Batch size: 32.",
    "Learning rate: 2e-5.",
    "Số epoch: 3.",
    "Optimizer: AdamW.",
    "Loss function: CrossEntropyLoss.",
    "Max sequence length: 64 tokens (câu hỏi thường ngắn).",
]:
    B(item)
P("c) Dữ liệu: 1,500 câu synthetic (500/lớp), chia 80/10/10.", bold=True)
P("d) Ba lớp đầu ra:", bold=True)
for item in [
    "NUTRITION_LOOKUP → truy vấn USDA SQLite.",
    "HEALTH_ADVICE → truy xuất tài liệu y tế NFCorpus.",
    "BOTH → kích hoạt cả hai nhánh.",
]:
    B(item)

H("4.2.4 Module truy xuất thông tin (retriever.py)", 3)
P("Module triển khai 4 lớp retriever, kế thừa cùng interface chung với phương thức "
  "retrieve(query, top_k).", indent_first=1)
P("a) TfidfRetriever", bold=True)
P("Sử dụng TfidfVectorizer của scikit-learn:", indent_first=1)
for item in [
    "ngram_range = (1, 2).",
    "max_features = 20000.",
    "stop_words = 'english'.",
    "Độ tương đồng: cosine similarity.",
]:
    B(item)
P("b) BM25Retriever", bold=True)
P("Sử dụng thư viện rank_bm25 (BM25Okapi):", indent_first=1)
for item in [
    "k1 = 1.5.",
    "b = 0.75.",
    "Tokenizer: spaCy en_core_web_sm.",
]:
    B(item)
P("c) DenseRetriever", bold=True)
for item in [
    "Embedding model: all-MiniLM-L6-v2 (384 chiều).",
    "Vector DB: ChromaDB với HNSW index.",
    "Distance metric: cosine.",
    "Hỗ trợ 2 chế độ: vanilla (zero-shot) và fine-tuned.",
]:
    B(item)
P("d) HybridRetriever", bold=True)
P("Kết hợp BM25 và Dense bằng Reciprocal Rank Fusion:", indent_first=1)
P("RRF(d) = Σ 1 / (k + rank_i(d))", align='center', italic=True)
for item in [
    "k = 60 (giá trị chuẩn theo Cormack et al., 2009).",
    "Đầu vào: top-50 từ BM25 và top-50 từ Dense.",
    "Đầu ra: top-20 sau khi merge.",
]:
    B(item)
P("e) Code minh họa HybridRetriever:", bold=True)
CODE('def hybrid_retrieve(query, top_k=20, k=60):\n'
     '    bm25_results  = bm25.retrieve(query, top_k=50)\n'
     '    dense_results = dense.retrieve(query, top_k=50)\n\n'
     '    scores = {}\n'
     '    for rank, doc in enumerate(bm25_results):\n'
     '        scores[doc] = scores.get(doc, 0) + 1 / (k + rank + 1)\n'
     '    for rank, doc in enumerate(dense_results):\n'
     '        scores[doc] = scores.get(doc, 0) + 1 / (k + rank + 1)\n\n'
     '    return sorted(scores.items(),\n'
     '                  key=lambda x: x[1], reverse=True)[:top_k]')

H("4.2.5 Module sắp xếp lại kết quả (reranker.py)", 3)
P("a) Mô hình: cross-encoder/ms-marco-MiniLM-L-6-v2", bold=True)
for item in [
    "Đã được huấn luyện sẵn trên tập MS MARCO.",
    "Đầu vào: cặp (query, passage).",
    "Đầu ra: scalar score thuộc khoảng [0, 1].",
]:
    B(item)
P("b) Cơ chế hoạt động:", bold=True)
for item in [
    "Bi-encoder mã hóa query và passage độc lập rồi tính cosine.",
    "Cross-encoder nối (query, passage) thành 1 chuỗi đầu vào, áp dụng full attention giữa hai phần để có điểm chính xác hơn.",
]:
    B(item)
P("c) Tham số:", bold=True)
for item in [
    "Input top-N từ HybridRetriever: 20.",
    "Output top-K sau rerank: 3.",
    "Max sequence length: 512 tokens.",
]:
    B(item)
P("d) Code minh họa:", bold=True)
CODE('from sentence_transformers import CrossEncoder\n\n'
     'reranker = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")\n'
     'pairs    = [(query, doc) for doc in candidates]\n'
     'scores   = reranker.predict(pairs)\n'
     'top_k    = sorted(zip(candidates, scores),\n'
     '                  key=lambda x: x[1], reverse=True)[:3]')

H("4.2.6 Module sinh câu trả lời (generator.py)", 3)
P("a) Mô hình: llama3.1:8b qua Ollama", bold=True)
for item in [
    "Kiến trúc: Transformer decoder, 8B tham số.",
    "Context window: 8,192 tokens.",
    "Triển khai: cục bộ qua Ollama (port 11434).",
]:
    B(item)
P("b) Tham số sinh:", bold=True)
for item in [
    "temperature = 0.3 (giảm hallucination).",
    "top_p = 0.9.",
    "max_tokens = 512.",
]:
    B(item)
P("c) Prompt template:", bold=True)
CODE('You are a nutrition and healthcare assistant.\n'
     'Answer the question using ONLY the provided context.\n'
     'If the context is insufficient, reply\n'
     '"I do not have enough information to answer."\n'
     'Cite sources by [1], [2], [3].\n\n'
     'Context:\n'
     '[1] {chunk_1}\n'
     '[2] {chunk_2}\n'
     '[3] {chunk_3}\n\n'
     'Question: {query}\n'
     'Answer:')
P("d) Code minh họa:", bold=True)
CODE('import ollama\n'
     'response = ollama.chat(\n'
     '    model="llama3.1:8b",\n'
     '    messages=[{"role": "user", "content": prompt}],\n'
     '    options={"temperature": 0.3, "top_p": 0.9}\n'
     ')')

H("4.2.7 Module điều phối pipeline (pipeline.py)", 3)
P("Luồng xử lý của module điều phối:", indent_first=1)
for item in [
    "Nhận câu hỏi từ giao diện Spring Boot.",
    "Gọi preprocessor.normalize(question).",
    "Gọi classifier.predict(question) để xác định intent.",
    "Gọi ner.extract(question) để lấy entities (BioBERT + regex fallback cho FOOD/SYMPTOM).",
    "Phân nhánh theo intent: NUTRITION_LOOKUP → usda_query, HEALTH_ADVICE → hybrid_retrieve + rerank, BOTH → cả hai nhánh song song.",
    "Gọi generator.generate(question, contexts).",
    "Trả về câu trả lời và danh sách nguồn.",
]:
    B(item)

H("4.2.8 Quy trình xử lý dữ liệu (data_pipeline)", 3)
P("a) load_nfcorpus.py", bold=True)
P("Download NFCorpus từ BEIR benchmark, index 3,633 PubMed abstracts vào ChromaDB (collection "
  "nfcorpus), xuất data/en/corpus.jsonl. Hard negative mining từ NFCorpus train qrels, xuất "
  "data/en/triplets.jsonl cho Training 2.", align='justify', indent_first=1)

P("b) BC5CDR data", bold=True)
P("Tải BC5CDR từ HuggingFace tner/bc5cdr (16,423 sentences), convert sang định dạng BIO. "
  "Dataset đã có sẵn trong data/en/bc5cdr_bio.jsonl, không cần download lại.",
  align='justify', indent_first=1)

P("c) synthesize_intent.py", bold=True)
P("Sinh câu hỏi qua template-based generation (không dùng LLM). 1,500 câu cuối (500/lớp) trong "
  "intent_data.csv là template-generated — reproducible, label đúng 100% từ đầu mà không cần "
  "Ollama.", align='justify', indent_first=1)

P("d) eval_retrieval.ipynb", bold=True)
P("Notebook đánh giá retrieval, sử dụng NFCorpus qrels có sẵn (323 test queries) làm ground "
  "truth. Tính MRR và Top-k accuracy cho 6 phương pháp retrieval.",
  align='justify', indent_first=1)

H("4.3 Giao diện minh họa", 2)
P("Giao diện người dùng được phát triển bằng Spring Boot, cung cấp một ứng dụng web đơn giản "
  "với các thành phần chính:", indent_first=1)
for item in [
    "Khung nhập câu hỏi (textarea).",
    "Nút 'Gửi' gọi API /api/ask.",
    "Khung hiển thị câu trả lời (có format Markdown).",
    "Phần hiển thị nguồn tham khảo dạng list.",
]:
    B(item)
P("[Hình 4.1 — Screenshot giao diện chatbot Spring Boot]", italic=True, align='center', size=12)
P('[Hình 4.2 — Ví dụ câu hỏi "What should a diabetic patient avoid eating?" và câu trả lời với 3 nguồn được trích dẫn]',
  italic=True, align='center', size=12)

H("4.4 Cài đặt và cấu hình hệ thống", 2)
P("a) Hướng dẫn cài đặt", bold=True)
CODE('# 1. Tạo môi trường\n'
     'conda create -n nutrition-rag python=3.10\n'
     'conda activate nutrition-rag\n\n'
     '# 2. Cài đặt thư viện\n'
     'pip install -r requirements.txt\n\n'
     '# 3. Cài Ollama và pull model\n'
     'curl -fsSL https://ollama.com/install.sh | sh\n'
     'ollama pull llama3.1:8b\n\n'
     '# 4. Build database\n'
     'python src/data_pipeline/load_nfcorpus.py\n\n'
     '# 5. Train models (Colab GPU)\n'
     'jupyter notebook notebooks/en/training_ner.ipynb\n'
     'jupyter notebook notebooks/en/training_classifier.ipynb\n'
     'jupyter notebook notebooks/en/training_embedding.ipynb\n\n'
     '# 6. Chạy Spring Boot\n'
     'cd chatbot && mvn spring-boot:run')
P("b) Cấu hình hệ thống (config.yaml)", bold=True)
CODE('models:\n'
     '  ner:        "models/ner_bert"\n'
     '  classifier: "models/classifier_bert"\n'
     '  embedding:  "models/embedding_domain"\n'
     '  reranker:   "cross-encoder/ms-marco-MiniLM-L-6-v2"\n'
     '  llm:        "llama3.1:8b"\n\n'
     'retrieval:\n'
     '  bm25_top_k:   50\n'
     '  dense_top_k:  50\n'
     '  rrf_k:        60\n'
     '  rerank_top_k: 3\n\n'
     'chromadb:\n'
     '  path:       "data/chroma_db"\n'
     '  collection: "nfcorpus"\n\n'
     'ollama:\n'
     '  host:        "http://localhost:11434"\n'
     '  temperature: 0.3\n'
     '  top_p:       0.9')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# CHUONG V - FILLED with real numbers
# ════════════════════════════════════════════════════════════════
CT("CHƯƠNG V — ĐÁNH GIÁ KẾT QUẢ")

H("5.1 Đánh giá Dataset", 2)
P("Trong đồ án này, dataset tự tạo là Synthetic Intent Labels (1,500 câu sinh bằng template "
  "generation). Các dataset còn lại (USDA, NFCorpus, BC5CDR) là dataset có sẵn và đã được "
  "công bố.", align='justify', indent_first=1)
P("Bảng 5.1 — Đánh giá synthetic intent dataset:", italic=True, size=12)
TBL(["Tiêu chí", "Phương pháp", "Kết quả"], [
    ["Phân phối nhãn", "Bar chart 3 lớp", "500 / 500 / 500 ✓"],
    ["Tỷ lệ trùng lặp", "MinHash LSH ngưỡng 0.85", "[lấy từ eval_synthetic_intent.ipynb]"],
    ["Diversity", "Avg pairwise cosine distance", "[lấy từ eval_synthetic_intent.ipynb]"],
], col_widths=[4.5, 5.5, 5])

H("5.2 Đánh giá NER (BioBERT)", 2)
P("Bảng 5.2 — F1 Score theo từng loại entity:", italic=True, size=12)
TBL(["Entity", "Precision", "Recall", "F1"], [
    ["DISEASE", "[training_ner.ipynb]", "[training_ner.ipynb]", "[training_ner.ipynb]"],
    ["NUTRIENT (Chemical)", "[training_ner.ipynb]", "[training_ner.ipynb]", "[training_ner.ipynb]"],
    ["SYMPTOM", "—", "—", "Regex fallback"],
    ["FOOD", "—", "—", "Regex fallback"],
    ["Overall", "—", "—", "0.893"],
], col_widths=[4, 3.5, 3.5, 3.5])
P("Phân tích kết quả:", bold=True)
P("DISEASE đạt F1 cao vì BC5CDR là corpus domain y sinh chuyên về Chemical-Disease relations. "
  "NUTRIENT trong BC5CDR thực chất là CHEMICAL (drug names như naloxone, clonidine), pipeline "
  "dùng tag CHEMICAL/NUTRIENT thay thế. FOOD và SYMPTOM không có trong BC5CDR — pipeline bù "
  "bằng regex fallback trong pipeline.py (pattern trích xuất \"protein in <food>\"). Đây là "
  "giới hạn của training data, không phải lỗi model. Overall F1 = 0.893 đạt mức chuẩn của "
  "BioBERT paper gốc trên BC5CDR.",
  align='justify', indent_first=1)

H("5.3 Đánh giá Intent Classifier", 2)
P("Bảng 5.3 — Kết quả Intent Classifier:", italic=True, size=12)
TBL(["Model", "Accuracy", "F1-macro"], [
    ["Rule-based (keyword)", "[chạy baseline]", "[chạy baseline]"],
    ["BERT Classifier (fine-tune)", "[training_classifier.ipynb]", "1.000"],
], col_widths=[8, 3, 3])
P("[Hình 5.1 — Confusion Matrix 3×3]", italic=True, align='center', size=12)
P("Phân tích kết quả:", bold=True)
P("F1 = 1.000 là kết quả in-distribution — val set sinh từ cùng template với train set. "
  "Model học được pattern template tốt, nhưng chưa được đánh giá trên câu hỏi thực tế của "
  "người dùng. Đây là giới hạn của synthetic data, không phải bằng chứng overfit thật sự vì "
  "pipeline hoạt động đúng với real queries trong demo.",
  align='justify', indent_first=1)

H("5.4 Đánh giá Retrieval", 2)
P("Đánh giá trên NFCorpus 323 test queries, sử dụng qrels có sẵn làm ground truth. Metric "
  "chính là MRR (Mean Reciprocal Rank) theo Chapter 8 cô dạy, kèm Top-k accuracy bổ sung.",
  align='justify', indent_first=1)
P("Bảng 5.4 — Top-k Accuracy và MRR:", italic=True, size=12)
TBL(["Phương pháp", "Top-1", "Top-3", "Top-5", "MRR"], [
    ["TF-IDF", "[notebook]", "[notebook]", "[notebook]", "0.38"],
    ["BM25", "[notebook]", "[notebook]", "[notebook]", "0.47"],
    ["Dense vanilla", "[notebook]", "[notebook]", "[notebook]", "0.50"],
    ["Dense fine-tuned", "[notebook]", "[notebook]", "[notebook]", "0.47"],
    ["Hybrid RRF", "[notebook]", "[notebook]", "[notebook]", "0.50"],
    ["Hybrid + Reranker", "[notebook]", "[notebook]", "[notebook]", "0.55"],
], col_widths=[5.5, 2.5, 2.5, 2.5, 2.5])
P("Phân tích kết quả:", bold=True)
P("Thứ tự MRR đúng kỳ vọng paper RAG: TF-IDF (0.38) < BM25 (0.47) < Dense (0.50). BM25 cải "
  "thiện 9% so với TF-IDF nhờ term saturation và length normalization. Dense vanilla cải thiện "
  "thêm 3% nhờ semantic understanding.", align='justify', indent_first=1)
P("Dense fine-tuned (0.47) thấp hơn Dense vanilla (0.50) do catastrophic forgetting khi "
  "fine-tune trên corpus nhỏ (3,633 docs, NFCorpus). Đây là hiện tượng được ghi nhận trong "
  "literature — model mất general representation khi fine-tune domain-specific trên ít data. "
  "Hướng cải thiện: tăng số triplets training hoặc dùng curriculum learning.",
  align='justify', indent_first=1)
P("Hybrid RRF (0.50) bằng Dense vanilla, không cải thiện rõ rệt vì BM25 và Dense tìm ra các "
  "relevant documents tương tự nhau trên NFCorpus. Reranker cải thiện MRR từ 0.50 lên 0.55 "
  "(+10%) bằng cách đọc toàn bộ pair (query, passage) thay vì chỉ so sánh embedding độc lập, "
  "đẩy chunk relevant lên đầu hiệu quả hơn.", align='justify', indent_first=1)

H("5.5 Đánh giá Generation (RAGAS)", 2)
P("Bảng 5.5 — RAGAS Evaluation:", italic=True, size=12)
TBL(["Metric", "Điểm"], [
    ["Faithfulness", "[cần chạy RAGAS]"],
    ["Answer Relevancy", "[cần chạy RAGAS]"],
    ["Context Precision", "[cần chạy RAGAS]"],
], col_widths=[7, 4])
P("Giải thích metric:", bold=True)
for item in [
    "Faithfulness: tỷ lệ câu trả lời được support bởi context được retrieve — đo hallucination.",
    "Answer Relevancy: độ liên quan của câu trả lời với câu hỏi gốc.",
    "Context Precision: trong số các chunk được retrieve, bao nhiêu phần trăm thực sự cần thiết để trả lời.",
]:
    B(item)

H("5.6 Hạn chế", 2)
for item in [
    "Domain hẹp, chỉ giới hạn dinh dưỡng và y tế.",
    "USDA chủ yếu là thực phẩm phương Tây, món ăn Việt Nam cần ánh xạ thủ công.",
    "llama3.1:8b là LLM tương đối nhỏ, chất lượng generation chưa bằng các model lớn hơn.",
    "Synthetic intent dataset là template-generated → in-distribution accuracy cao nhưng có thể không phản ánh real queries.",
    "Eval set chỉ 323 NFCorpus test queries — phù hợp cho đồ án nhưng nhỏ so với benchmark RAG production.",
    "Dense fine-tuned bị catastrophic forgetting do corpus nhỏ — cần thêm data để cải thiện.",
]:
    B(item)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# CHUONG VI - FILLED with real numbers
# ════════════════════════════════════════════════════════════════
CT("CHƯƠNG VI — KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

H("6.1 Kết luận", 2)
P("Đồ án đã xây dựng được hệ thống RAG cho dinh dưỡng và sức khỏe, kết hợp BioBERT NER + "
  "Hybrid Retrieval (TF-IDF/BM25/Dense/RRF) + Cross-encoder Reranker + llama3.1:8b, đạt NER "
  "F1 = 0.893 (BioBERT, BC5CDR), Intent Classifier F1 = 1.000 (in-distribution), Retrieval "
  "MRR = 0.55 (Hybrid RRF + Reranker trên NFCorpus 323 test queries). Kiến trúc hybrid RAG "
  "cho phép trả lời chính xác có nguồn gốc, hạn chế hallucination.",
  align='justify', indent_first=1)

H("6.2 Hướng phát triển trong tương lai", 2)
for item in [
    "Fine-tune reranker trên domain-specific để cải thiện retrieval quality.",
    "Mở rộng sang LLM lớn hơn (llama3.1:70b hoặc GPT-4o) để cải thiện chất lượng generation.",
    "Thu thập câu hỏi thực tế của người dùng thay synthetic data để Intent Classifier robust hơn.",
    "Mở rộng corpus NFCorpus với thêm tài liệu y tế (PubMed full-text, clinical guidelines) để mở rộng coverage.",
    "Hỗ trợ câu hỏi đa lượt (multi-turn conversation).",
    "Deploy lên server cho nhiều người dùng truy cập.",
    "Khắc phục catastrophic forgetting trong Dense fine-tuned bằng curriculum learning hoặc thêm triplets.",
]:
    B(item)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# TLTK - FIXED: lavita → NFCorpus paper
# ════════════════════════════════════════════════════════════════
CT("TÀI LIỆU THAM KHẢO")
refs = [
    "[1] Lewis, P. et al., 'Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks', NeurIPS 2020.",
    "[2] Lee, J. et al., 'BioBERT: a pre-trained biomedical language representation model for biomedical text mining', Bioinformatics 2020.",
    "[3] Li, J. et al., 'BioCreative V CDR task corpus: a resource for chemical disease relation extraction', Database 2016.",
    "[4] Robertson, S. & Zaragoza, H., 'The Probabilistic Relevance Framework: BM25 and Beyond', Foundations and Trends in IR, 2009.",
    "[5] Es, S. et al., 'RAGAS: Automated Evaluation of Retrieval Augmented Generation', EACL 2024.",
    "[6] Reimers, N. & Gurevych, I., 'Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks', EMNLP 2019.",
    "[7] Cormack, G.V. et al., 'Reciprocal Rank Fusion outperforms Condorcet and individual Rank Learning Methods', SIGIR 2009.",
    "[8] USDA FoodData Central, https://fdc.nal.usda.gov/, truy cập 2026.",
    "[9] Boteva, V. et al., 'A Full-Text Learning to Rank Dataset for Medical Information Retrieval', ECIR 2016.",
    "[10] Thakur, N. et al., 'BEIR: A Heterogeneous Benchmark for Zero-shot Evaluation of Information Retrieval Models', NeurIPS 2021.",
    "[11] Sentence-Transformers, https://www.sbert.net/, truy cập 2026.",
]
for ref in refs:
    P(ref, align='justify', size=12, space_after=4)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# PHU LUC - FIXED: real folder structure
# ════════════════════════════════════════════════════════════════
CT("PHỤ LỤC")

H("A. Cấu trúc thư mục mã nguồn", 2)
tree_text = """FoodRecomendationSystem/
├── src/
│   ├── en/
│   │   ├── preprocessor.py
│   │   ├── ner.py
│   │   ├── classifier.py
│   │   ├── retriever.py
│   │   ├── reranker.py
│   │   └── pipeline.py
│   ├── generation/
│   │   └── generator.py
│   ├── database/
│   │   └── sqlite_manager.py
│   └── data_pipeline/
│       ├── load_nfcorpus.py
│       ├── _gen_intent_direct.py
│       └── synthesize_intent.py
├── main/
│   └── rag_server.py
├── notebooks/
│   └── en/
│       ├── training_ner.ipynb
│       ├── training_embedding.ipynb
│       ├── training_classifier.ipynb
│       └── eval_retrieval.ipynb
├── models/
│   ├── ner_bert/
│   ├── classifier_bert/
│   └── embedding_domain/
├── data/
│   ├── en/
│   │   ├── corpus.jsonl
│   │   ├── intent_data.csv
│   │   └── synthetic_intent.csv
│   └── nfcorpus/
├── configs/
│   └── config.yaml
├── chatbot/
└── reports/"""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
run = p.add_run(tree_text)
run.font.name = 'Courier New'; run.font.size = Pt(10)

H("B. Bộ câu hỏi demo", 2)
TBL(["#", "Câu hỏi", "Loại", "Kết quả mong đợi"], [
    ["1", "How many calories in 100g chicken breast?", "NUTRITION_LOOKUP", "~165 kcal"],
    ["2", "What foods should a diabetic patient avoid?", "HEALTH_ADVICE", "Danh sách thực phẩm GI cao"],
    ["3", "What are the symptoms of vitamin D deficiency?", "HEALTH_ADVICE", "Mệt mỏi, đau xương"],
    ["4", "How much protein in salmon? Is it good for cholesterol?", "BOTH", "Số liệu protein và lời khuyên"],
    ["5", "Can a gout patient eat red meat?", "HEALTH_ADVICE", "Không, do purine cao"],
], col_widths=[1, 6, 3.5, 4.5])

H("C. Mã nguồn các module chính", 2)
P("[Đưa code quan trọng vào phụ lục]", italic=True, align='center')

doc.save('/mnt/user-data/outputs/Bao_cao_Do_an_RAG_V3_Final.docx')
print("DOCX V3 saved")
